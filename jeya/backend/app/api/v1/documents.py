import uuid
import io
from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc
from typing import List
from app.core.database import get_db
from app.models.document import Document
from app.models.user import User
from app.schemas.document import DocumentCreate, DocumentGenerate, DocumentResponse
from app.dependencies.auth import get_current_user
from app.services.llm_service import generate_document

router = APIRouter()


@router.get("", response_model=List[DocumentResponse])
async def list_documents(current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.user_id == current_user.id).order_by(desc(Document.created_at)))
    return [DocumentResponse.model_validate(d) for d in result.scalars().all()]


@router.post("", response_model=DocumentResponse)
async def create_document(data: DocumentCreate, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    doc = Document(user_id=current_user.id, title=data.title, doc_type=data.doc_type, content=data.content or "")
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return DocumentResponse.model_validate(doc)


@router.post("/generate", response_model=DocumentResponse)
async def generate_doc(data: DocumentGenerate, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    content = await generate_document(data.title, data.doc_type, data.prompt)
    doc = Document(user_id=current_user.id, title=data.title, doc_type=data.doc_type, content=content)
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return DocumentResponse.model_validate(doc)


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: uuid.UUID, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == document_id, Document.user_id == current_user.id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    return DocumentResponse.model_validate(doc)


@router.delete("/{document_id}")
async def delete_document(document_id: uuid.UUID, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == document_id, Document.user_id == current_user.id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    await db.delete(doc)
    await db.commit()
    return {"message": "Hujjat o'chirildi"}


@router.get("/{document_id}/export")
async def export_document(document_id: uuid.UUID, format: str = "pdf", current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == document_id, Document.user_id == current_user.id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if format == "docx":
        from docx import Document as DocxDocument
        docx = DocxDocument()
        docx.add_heading(doc.title, 0)
        for para in doc.content.split("\n"):
            if para.strip():
                docx.add_paragraph(para)
        buf = io.BytesIO()
        docx.save(buf)
        buf.seek(0)
        return Response(content=buf.read(),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={"Content-Disposition": f'attachment; filename="{doc.title}.docx"'})
    else:
        content_bytes = f"{doc.title}\n\n{doc.content}".encode("utf-8")
        return Response(content=content_bytes, media_type="text/plain",
                        headers={"Content-Disposition": f'attachment; filename="{doc.title}.txt"'})
